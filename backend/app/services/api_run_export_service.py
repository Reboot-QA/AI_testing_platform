import glob
import json
import os
import platform
from datetime import datetime
from io import BytesIO
from typing import Any, Iterable, List, Optional, Tuple
from urllib.parse import quote
from xml.sax.saxutils import escape

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from app.schemas import ApiTestRunDetailOut

STATUS_LABELS = {
    "passed": "通过",
    "failed": "失败",
    "running": "执行中",
    "skipped": "跳过",
}

EXPORT_MEDIA_TYPES = {
    "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "json": "application/json; charset=utf-8",
}

EXPORT_EXTENSIONS = {
    "excel": "xlsx",
    "word": "docx",
    "pdf": "pdf",
    "json": "json",
}

SUPPORTED_EXPORT_FORMATS = set(EXPORT_MEDIA_TYPES)


def build_content_disposition(filename: str, fallback: str = "report.xlsx") -> str:
    encoded = quote(filename, safe="")
    return f'attachment; filename="{fallback}"; filename*=UTF-8\'\'{encoded}'


def build_export_filename(reports: List[ApiTestRunDetailOut], ext: str) -> str:
    if len(reports) == 1:
        report = reports[0]
        suite = (report.suite_name or "report").replace("/", "_").replace("\\", "_")
        return f"测试报告_{report.id}_{suite}.{ext}"
    return f"测试报告_批量导出_{len(reports)}条.{ext}"


def build_run_export(
    reports: List[ApiTestRunDetailOut], export_format: str
) -> Tuple[Any, str, str]:
    normalized = (export_format or "excel").lower()
    if normalized not in SUPPORTED_EXPORT_FORMATS:
        raise ValueError("不支持的导出格式")

    if normalized == "excel":
        return build_run_export_excel(reports), EXPORT_MEDIA_TYPES["excel"], EXPORT_EXTENSIONS["excel"]
    if normalized == "word":
        return build_run_export_word(reports), EXPORT_MEDIA_TYPES["word"], EXPORT_EXTENSIONS["word"]
    if normalized == "pdf":
        _ensure_reportlab()
        return build_run_export_pdf(reports), EXPORT_MEDIA_TYPES["pdf"], EXPORT_EXTENSIONS["pdf"]
    content = build_run_export_json(reports, single=len(reports) == 1)
    return content, EXPORT_MEDIA_TYPES["json"], EXPORT_EXTENSIONS["json"]


def _format_dt(value: Optional[datetime]) -> str:
    if not value:
        return ""
    return value.strftime("%Y-%m-%d %H:%M:%S")


def _format_duration(ms: float) -> str:
    if ms is None:
        return ""
    if ms < 1000:
        return f"{int(round(ms))}ms"
    return f"{ms / 1000:.2f}s"


def _pretty_json_text(text: Optional[str]) -> str:
    raw = str(text or "").strip()
    if not raw:
        return "-"
    try:
        return json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        return raw


def _format_assertions(results: Iterable[Any]) -> str:
    lines = []
    for item in results:
        if hasattr(item, "model_dump"):
            data = item.model_dump()
        elif isinstance(item, dict):
            data = item
        else:
            continue
        status = "通过" if data.get("passed") else "失败"
        message = data.get("message") or data.get("type") or ""
        expected = data.get("expected")
        actual = data.get("actual")
        lines.append(f"[{status}] {message} | 期望: {expected} | 实际: {actual}")
    return "\n".join(lines) if lines else "-"


def _report_summary_rows(report: ApiTestRunDetailOut) -> List[List[str]]:
    return [
        ["报告 ID", str(report.id)],
        ["测试套件", report.suite_name or "-"],
        ["执行结果", STATUS_LABELS.get(report.status, report.status)],
        ["通过率", f"{report.pass_rate}%"],
        ["通过 / 失败 / 总数", f"{report.passed_count} / {report.failed_count} / {report.total_count}"],
        ["总耗时", _format_duration(report.duration_ms)],
        ["触发方式", report.triggered_by or "-"],
        ["开始时间", _format_dt(report.started_at)],
        ["结束时间", _format_dt(report.finished_at)],
    ]


def _assertion_table_rows(step) -> List[List[str]]:
    rows = []
    for item in step.assertion_results or []:
        if hasattr(item, "model_dump"):
            data = item.model_dump()
        elif isinstance(item, dict):
            data = item
        else:
            continue
        rows.append(
            [
                str(data.get("type") or ""),
                str(data.get("message") or ""),
                str(data.get("expected") if data.get("expected") is not None else "-"),
                str(data.get("actual") if data.get("actual") is not None else "-"),
                "通过" if data.get("passed") else "失败",
            ]
        )
    return rows


def _autosize_row(ws, row_num: int, min_height: int = 18, max_height: int = 120) -> None:
    max_lines = 1
    for cell in ws[row_num]:
        if cell.value is None:
            continue
        text = str(cell.value)
        col_width = ws.column_dimensions[get_column_letter(cell.column)].width or 10
        chars_per_line = max(int(col_width * 1.2), 10)
        lines = 0
        for part in text.split("\n"):
            lines += max(1, (len(part) + chars_per_line - 1) // chars_per_line)
        max_lines = max(max_lines, lines)
    ws.row_dimensions[row_num].height = min(max(min_height, max_lines * 15), max_height)


def _configure_sheet(ws, column_widths: dict) -> None:
    header_font = Font(bold=True)
    for header in ws[1]:
        header.font = header_font
        header.alignment = Alignment(horizontal="center", vertical="center")
    for column, width in column_widths.items():
        ws.column_dimensions[column].width = width
    ws.row_dimensions[1].height = 20


def _apply_multiline_cell(ws, row_num: int, column: int, value: Any) -> None:
    cell = ws.cell(row_num, column, value if value is not None else "")
    cell.alignment = Alignment(vertical="top", wrap_text=True)


def build_run_export_excel(reports: List[ApiTestRunDetailOut]) -> BytesIO:
    wb = Workbook()

    summary_ws = wb.active
    summary_ws.title = "报告摘要"
    summary_headers = [
        "报告 ID",
        "套件",
        "结果",
        "通过率",
        "通过数",
        "失败数",
        "总数",
        "耗时",
        "触发方式",
        "开始时间",
        "结束时间",
    ]
    summary_ws.append(summary_headers)
    _configure_sheet(
        summary_ws,
        {"A": 10, "B": 24, "C": 10, "D": 10, "E": 8, "F": 8, "G": 8, "H": 10, "I": 12, "J": 20, "K": 20},
    )

    for report in reports:
        row_num = summary_ws.max_row + 1
        summary_ws.cell(row_num, 1, report.id)
        summary_ws.cell(row_num, 2, report.suite_name)
        summary_ws.cell(row_num, 3, STATUS_LABELS.get(report.status, report.status))
        summary_ws.cell(row_num, 4, f"{report.pass_rate}%")
        summary_ws.cell(row_num, 5, report.passed_count)
        summary_ws.cell(row_num, 6, report.failed_count)
        summary_ws.cell(row_num, 7, report.total_count)
        summary_ws.cell(row_num, 8, _format_duration(report.duration_ms))
        summary_ws.cell(row_num, 9, report.triggered_by or "")
        summary_ws.cell(row_num, 10, _format_dt(report.started_at))
        summary_ws.cell(row_num, 11, _format_dt(report.finished_at))

    detail_ws = wb.create_sheet("用例明细")
    detail_headers = [
        "报告 ID",
        "套件",
        "序号",
        "用例名称",
        "方法",
        "URL",
        "结果",
        "耗时",
        "响应状态",
        "错误信息",
        "断言结果",
        "请求头",
        "请求体",
        "响应头",
        "响应体",
    ]
    detail_ws.append(detail_headers)
    _configure_sheet(
        detail_ws,
        {
            "A": 10, "B": 20, "C": 8, "D": 22, "E": 10, "F": 36, "G": 10, "H": 10,
            "I": 10, "J": 24, "K": 28, "L": 24, "M": 24, "N": 24, "O": 32,
        },
    )

    for report in reports:
        for index, step in enumerate(report.step_results, start=1):
            row_num = detail_ws.max_row + 1
            detail_ws.cell(row_num, 1, report.id)
            detail_ws.cell(row_num, 2, report.suite_name)
            detail_ws.cell(row_num, 3, index)
            detail_ws.cell(row_num, 4, step.case_name)
            detail_ws.cell(row_num, 5, step.method)
            _apply_multiline_cell(detail_ws, row_num, 6, step.url)
            detail_ws.cell(row_num, 7, STATUS_LABELS.get(step.status, step.status))
            detail_ws.cell(row_num, 8, _format_duration(step.duration_ms))
            detail_ws.cell(row_num, 9, step.response_status if step.response_status is not None else "")
            _apply_multiline_cell(detail_ws, row_num, 10, step.error_message or "")
            _apply_multiline_cell(detail_ws, row_num, 11, _format_assertions(step.assertion_results))
            _apply_multiline_cell(detail_ws, row_num, 12, _pretty_json_text(step.request_headers))
            _apply_multiline_cell(detail_ws, row_num, 13, _pretty_json_text(step.request_body))
            _apply_multiline_cell(detail_ws, row_num, 14, _pretty_json_text(step.response_headers))
            _apply_multiline_cell(detail_ws, row_num, 15, _pretty_json_text(step.response_body))
            _autosize_row(detail_ws, row_num)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def build_run_export_json(reports: List[ApiTestRunDetailOut], *, single: bool = False) -> bytes:
    payload = reports[0].model_dump(mode="json") if single and len(reports) == 1 else [
        report.model_dump(mode="json") for report in reports
    ]
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _set_docx_font(run, name="微软雅黑", size=11, bold=False):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def _docx_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    size = 18 if level == 1 else 14 if level == 2 else 12
    for run in paragraph.runs:
        _set_docx_font(run, size=size, bold=True)
    return paragraph


def _docx_paragraph(doc, text: str, *, bold=False, monospace=False):
    from docx.shared import Cm, Pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font_name = "Consolas" if monospace else "微软雅黑"
    _set_docx_font(run, name=font_name, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if monospace:
        paragraph.paragraph_format.left_indent = Cm(0.4)
    return paragraph


def _docx_table(doc, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            _set_docx_font(run, bold=True, size=10)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index + 1].cells[col_index]
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                _set_docx_font(run, size=9)
    doc.add_paragraph()


def _docx_code_block(doc, title: str, content: Optional[str]):
    _docx_paragraph(doc, title, bold=True)
    _docx_paragraph(doc, _pretty_json_text(content), monospace=True)


def _append_report_word(doc, report: ApiTestRunDetailOut) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("接口自动化测试报告")
    _set_docx_font(run, size=20, bold=True)

    _docx_paragraph(doc, f"报告编号：{report.id}    套件：{report.suite_name or '-'}", bold=True)
    _docx_heading(doc, "报告摘要", level=2)
    _docx_table(doc, ["项目", "内容"], _report_summary_rows(report))

    _docx_heading(doc, "用例执行明细", level=2)
    for index, step in enumerate(report.step_results, start=1):
        status = STATUS_LABELS.get(step.status, step.status)
        _docx_heading(doc, f"{index}. {step.case_name}（{status}）", level=3)
        _docx_paragraph(
            doc,
            f"方法：{step.method}    耗时：{_format_duration(step.duration_ms)}    "
            f"响应状态：{step.response_status if step.response_status is not None else '-'}",
        )
        _docx_paragraph(doc, f"请求 URL：{step.url}")

        if step.error_message:
            _docx_paragraph(doc, f"错误信息：{step.error_message}", bold=True)

        _docx_code_block(doc, "请求头", step.request_headers)
        _docx_code_block(doc, "请求体", step.request_body)
        _docx_code_block(doc, "响应头", step.response_headers)
        _docx_code_block(doc, "响应体", step.response_body)

        assertion_rows = _assertion_table_rows(step)
        if assertion_rows:
            _docx_paragraph(doc, "断言结果", bold=True)
            _docx_table(doc, ["类型", "说明", "期望", "实际", "结果"], assertion_rows)


def build_run_export_word(reports: List[ApiTestRunDetailOut]) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for index, report in enumerate(reports):
        if index > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _append_report_word(doc, report)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _ensure_reportlab() -> None:
    try:
        import reportlab  # noqa: F401
    except ImportError as exc:
        raise RuntimeError("PDF 导出依赖 reportlab，请执行 pip install -r requirements.txt 并重启后端") from exc


def _pdf_font_candidates() -> List[str]:
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyhbd.ttc",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    if platform.system() == "Darwin":
        candidates.extend(
            [
                "/System/Library/Fonts/PingFang.ttc",
                "/Library/Fonts/Arial Unicode.ttf",
            ]
        )
    for pattern in (
        "/usr/share/fonts/**/NotoSansCJK*.ttc",
        "/usr/share/fonts/**/NotoSansSC*.otf",
        "/usr/share/fonts/**/wqy-microhei*.ttc",
    ):
        candidates.extend(glob.glob(pattern, recursive=True))
    seen = set()
    unique: List[str] = []
    for path in candidates:
        normalized = os.path.normpath(path)
        if normalized not in seen:
            seen.add(normalized)
            unique.append(normalized)
    return unique


def _register_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    registered = set(pdfmetrics.getRegisteredFontNames())
    if "ReportCJK" in registered:
        return "ReportCJK"
    if "STSong-Light" in registered:
        return "STSong-Light"

    for path in _pdf_font_candidates():
        if not os.path.exists(path):
            continue
        try:
            font_name = "ReportCJK"
            if path.lower().endswith(".ttc"):
                pdfmetrics.registerFont(TTFont(font_name, path, subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont(font_name, path))
            return font_name
        except Exception:
            continue

    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def _pdf_code_text(text: Optional[str], limit: int = 12000) -> str:
    pretty = _pretty_json_text(text)
    if len(pretty) <= limit:
        return pretty
    return f"{pretty[:limit]}\n... (内容过长，PDF 中已截断)"


def _pdf_escape(text: Any) -> str:
    raw = str(text if text is not None else "-").replace("\x00", "")
    return escape(raw).replace("\n", "<br/>")


def _build_pdf_styles(font_name: str):
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "heading2": ParagraphStyle(
            "ReportHeading2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=8,
            spaceAfter=8,
        ),
        "heading3": ParagraphStyle(
            "ReportHeading3",
            parent=base["Heading3"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            spaceBefore=6,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "ReportBody",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "code": ParagraphStyle(
            "ReportCode",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=8,
            leading=11,
            leftIndent=12,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "table": ParagraphStyle(
            "ReportTable",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=9,
            leading=12,
            wordWrap="CJK",
        ),
    }


def _pdf_table(rows: List[List[str]], col_widths: List[float], styles) -> Any:
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle

    wrapped = [
        [Paragraph(_pdf_escape(cell), styles["table"]) for cell in row]
        for row in rows
    ]
    table = Table(wrapped, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2ff")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("FONTNAME", (0, 0), (-1, -1), styles["table"].fontName),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _append_report_pdf(story: List[Any], report: ApiTestRunDetailOut, styles) -> None:
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, Spacer

    story.append(Paragraph("接口自动化测试报告", styles["title"]))
    story.append(
        Paragraph(
            _pdf_escape(f"报告编号：{report.id}    套件：{report.suite_name or '-'}"),
            styles["body"],
        )
    )
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("报告摘要", styles["heading2"]))
    story.append(
        _pdf_table(
            [["项目", "内容"], *_report_summary_rows(report)],
            [4.2 * cm, 12.5 * cm],
            styles,
        )
    )
    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("用例执行明细", styles["heading2"]))

    for index, step in enumerate(report.step_results, start=1):
        status = STATUS_LABELS.get(step.status, step.status)
        story.append(Paragraph(_pdf_escape(f"{index}. {step.case_name}（{status}）"), styles["heading3"]))
        story.append(
            Paragraph(
                _pdf_escape(
                    f"方法：{step.method}    耗时：{_format_duration(step.duration_ms)}    "
                    f"响应状态：{step.response_status if step.response_status is not None else '-'}"
                ),
                styles["body"],
            )
        )
        story.append(Paragraph(_pdf_escape(f"请求 URL：{step.url}"), styles["body"]))
        if step.error_message:
            story.append(Paragraph(_pdf_escape(f"错误信息：{step.error_message}"), styles["body"]))

        for title, content in [
            ("请求头", step.request_headers),
            ("请求体", step.request_body),
            ("响应头", step.response_headers),
            ("响应体", step.response_body),
        ]:
            story.append(Paragraph(_pdf_escape(title), styles["body"]))
            story.append(Paragraph(_pdf_escape(_pdf_code_text(content)), styles["code"]))

        assertion_rows = _assertion_table_rows(step)
        if assertion_rows:
            story.append(Paragraph("断言结果", styles["body"]))
            story.append(
                _pdf_table(
                    [["类型", "说明", "期望", "实际", "结果"], *assertion_rows],
                    [2.2 * cm, 4.5 * cm, 2.8 * cm, 2.8 * cm, 1.8 * cm],
                    styles,
                )
            )
        story.append(Spacer(1, 0.2 * cm))


def build_run_export_pdf(reports: List[ApiTestRunDetailOut]) -> BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import PageBreak, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    font_name = _register_pdf_font()
    styles = _build_pdf_styles(font_name)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=42,
        bottomMargin=36,
        title="接口自动化测试报告",
    )

    story: List[Any] = []
    for index, report in enumerate(reports):
        if index > 0:
            story.append(PageBreak())
        _append_report_pdf(story, report, styles)
    story.append(Spacer(1, 12))
    try:
        doc.build(story)
    except Exception as exc:
        raise RuntimeError(f"PDF 生成失败: {exc}") from exc
    buffer.seek(0)
    return buffer
