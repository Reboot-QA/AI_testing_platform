"""功能测试单（手工执行）报告导出 · Excel / Word / PDF。"""

from datetime import datetime
from io import BytesIO
from typing import Any, List, Optional, Tuple
from urllib.parse import quote

from app.schemas import ManualTestRunDetailOut

RUN_STATUS_LABELS = {
    "waiting": "待开始",
    "running": "执行中",
    "finished": "已完成",
}

CASE_RESULT_LABELS = {
    "pending": "待测",
    "pass": "通过",
    "fail": "失败",
    "blocked": "阻塞",
    "skip": "跳过",
}

EXPORT_MEDIA_TYPES = {
    "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

EXPORT_EXTENSIONS = {
    "excel": "xlsx",
    "word": "docx",
    "pdf": "pdf",
}

SUPPORTED_EXPORT_FORMATS = set(EXPORT_MEDIA_TYPES)


def _format_dt(value: Optional[datetime]) -> str:
    if not value:
        return "-"
    return value.strftime("%Y-%m-%d %H:%M:%S")


def _format_report_time(value: Optional[datetime]) -> str:
    if not value:
        return ""
    return value.strftime("%Y%m%d%H%M%S")


def _format_run_duration(started: Optional[datetime], finished: Optional[datetime]) -> str:
    if not started or not finished:
        return "-"
    total_sec = int(round((finished - started).total_seconds()))
    if total_sec < 0:
        return "-"
    if total_sec < 60:
        return f"{total_sec} 秒"
    minutes, seconds = divmod(total_sec, 60)
    return f"{minutes} 分 {seconds} 秒" if seconds else f"{minutes} 分"


def build_content_disposition(filename: str, fallback: str = "report.xlsx") -> str:
    encoded = quote(filename, safe="")
    return f'attachment; filename="{fallback}"; filename*=UTF-8\'\'{encoded}'


def build_export_filename(report: ManualTestRunDetailOut, ext: str) -> str:
    name = (report.name or "report").replace("/", "_").replace("\\", "_")
    time_token = _format_report_time(report.finished_at or report.started_at or report.created_at)
    return f"功能测试报告_{time_token}_{name}.{ext}"


def _summary_rows(report: ManualTestRunDetailOut) -> List[List[str]]:
    return [
        ["测试单", report.name or "-"],
        ["版本", report.build_name or "-"],
        ["状态", RUN_STATUS_LABELS.get(report.status, report.status)],
        ["通过率", f"{report.pass_rate}%"],
        [
            "结果统计",
            f"通过 {report.passed_count} / 失败 {report.failed_count} / "
            f"阻塞 {report.blocked_count} / 跳过 {report.skipped_count} / "
            f"待测 {report.pending_count} / 共 {report.total_count}",
        ],
        ["执行人", report.executor_name or "-"],
        ["耗时", _format_run_duration(report.started_at, report.finished_at)],
        ["开始时间", _format_dt(report.started_at)],
        ["结束时间", _format_dt(report.finished_at)],
        ["说明", report.description or "-"],
    ]


def _case_table_rows(report: ManualTestRunDetailOut) -> List[List[str]]:
    rows: List[List[str]] = []
    for case in report.cases:
        rows.append(
            [
                str(case.testcase_sort_order or case.sort_order or ""),
                case.case_title or "-",
                case.case_priority or "-",
                case.case_type or "-",
                CASE_RESULT_LABELS.get(case.result, case.result),
                case.executor_name or "-",
                case.actual_result or "-",
                case.remark or "-",
                case.preconditions or "-",
                case.steps or "-",
                case.expected_results or "-",
            ]
        )
    return rows


CASE_HEADERS = [
    "序号",
    "用例",
    "优先级",
    "类型",
    "结果",
    "执行人",
    "实际结果",
    "备注",
    "前置条件",
    "步骤",
    "期望结果",
]


def build_manual_run_export_excel(report: ManualTestRunDetailOut) -> BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    wb = Workbook()
    summary_ws = wb.active
    summary_ws.title = "报告摘要"
    summary_ws.append(["项目", "内容"])
    for cell in summary_ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row in _summary_rows(report):
        summary_ws.append(row)
    summary_ws.column_dimensions["A"].width = 16
    summary_ws.column_dimensions["B"].width = 56

    detail_ws = wb.create_sheet("用例明细")
    detail_ws.append(CASE_HEADERS)
    for cell in detail_ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in _case_table_rows(report):
        detail_ws.append(row)
        row_num = detail_ws.max_row
        for col in range(1, len(CASE_HEADERS) + 1):
            detail_ws.cell(row_num, col).alignment = Alignment(vertical="top", wrap_text=True)
    widths = {"A": 8, "B": 28, "C": 10, "D": 10, "E": 10, "F": 12, "G": 24, "H": 18, "I": 20, "J": 28, "K": 28}
    for col, width in widths.items():
        detail_ws.column_dimensions[col].width = width

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _set_docx_font(run, name="微软雅黑", size=11, bold=False):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def _docx_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    size = 18 if level == 1 else 14 if level == 2 else 12
    for run in paragraph.runs:
        _set_docx_font(run, size=size, bold=True)


def _docx_paragraph(doc, text: str, *, bold=False):
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_docx_font(run, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25


def _docx_table(doc, headers: List[str], rows: List[List[str]]):
    if not rows:
        _docx_paragraph(doc, "（无用例明细）")
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            _set_docx_font(run, bold=True, size=10)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index + 1].cells[col_index]
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                _set_docx_font(run, size=9)
    doc.add_paragraph()


def build_manual_run_export_word(report: ManualTestRunDetailOut) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("功能测试报告")
    _set_docx_font(run, size=20, bold=True)
    _docx_paragraph(doc, f"测试单：{report.name or '-'}    版本：{report.build_name or '-'}", bold=True)

    _docx_heading(doc, "报告摘要", level=2)
    _docx_table(doc, ["项目", "内容"], _summary_rows(report))

    _docx_heading(doc, "用例执行明细", level=2)
    _docx_table(doc, CASE_HEADERS, _case_table_rows(report))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_manual_run_export_pdf(report: ManualTestRunDetailOut) -> BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    from app.services.apifox.run_export_pdf import (
        _build_pdf_styles,
        _ensure_reportlab,
        _pdf_escape,
        _pdf_table,
        _register_pdf_font,
    )

    _ensure_reportlab()
    buffer = BytesIO()
    font_name = _register_pdf_font()
    styles = _build_pdf_styles(font_name)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=42,
        bottomMargin=36,
        title="功能测试报告",
    )

    story: List[Any] = []
    story.append(Paragraph("功能测试报告", styles["title"]))
    story.append(
        Paragraph(
            _pdf_escape(f"测试单：{report.name or '-'}    版本：{report.build_name or '-'}"),
            styles["body"],
        )
    )
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("报告摘要", styles["heading2"]))
    story.append(_pdf_table([["项目", "内容"], *_summary_rows(report)], [4.2 * cm, 12.5 * cm], styles))
    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("用例执行明细", styles["heading2"]))
    pdf_case_headers = ["序号", "用例", "优先级", "结果", "执行人", "实际结果", "备注"]
    pdf_case_rows = [
        row[:2] + row[2:3] + row[4:5] + row[5:8] for row in _case_table_rows(report)
    ]
    if pdf_case_rows:
        story.append(
            _pdf_table(
                [pdf_case_headers, *pdf_case_rows],
                [1.0 * cm, 4.8 * cm, 1.6 * cm, 1.4 * cm, 2.0 * cm, 3.6 * cm, 2.8 * cm],
                styles,
            )
        )
    else:
        story.append(Paragraph("（无用例明细）", styles["body"]))
    story.append(Spacer(1, 12))

    try:
        doc.build(story)
    except Exception as exc:
        raise RuntimeError(f"PDF 生成失败: {exc}") from exc
    buffer.seek(0)
    return buffer


def build_manual_run_export(
    report: ManualTestRunDetailOut, export_format: str
) -> Tuple[Any, str, str]:
    normalized = (export_format or "excel").lower()
    if normalized not in SUPPORTED_EXPORT_FORMATS:
        raise ValueError("不支持的导出格式")

    if normalized == "excel":
        return (
            build_manual_run_export_excel(report),
            EXPORT_MEDIA_TYPES["excel"],
            EXPORT_EXTENSIONS["excel"],
        )
    if normalized == "word":
        return (
            build_manual_run_export_word(report),
            EXPORT_MEDIA_TYPES["word"],
            EXPORT_EXTENSIONS["word"],
        )
    return (
        build_manual_run_export_pdf(report),
        EXPORT_MEDIA_TYPES["pdf"],
        EXPORT_EXTENSIONS["pdf"],
    )
