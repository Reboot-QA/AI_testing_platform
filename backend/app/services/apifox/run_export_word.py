"""Apifox 运行报告导出 · Word。"""

from io import BytesIO
from typing import Any, List

from app.routers.apifox.run_schemas import RunOut
from app.services.apifox.run_export_common import (
    STATUS_LABELS,
    _assertion_table_rows,
    _format_duration,
    _format_report_time,
    _pretty_json_text,
    _report_summary_rows,
)


def _set_docx_font(run, name="微软雅黑", size=11, bold=False):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def _docx_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    size = 18 if level == 1 else 14 if level == 2 else 12
    for run in paragraph.runs:
        _set_docx_font(run, size=size, bold=True)
    return paragraph


def _docx_paragraph(doc, text: str, *, bold=False, monospace=False):
    from docx.shared import Cm, Pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font_name = "Consolas" if monospace else "微软雅黑"
    _set_docx_font(run, name=font_name, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if monospace:
        paragraph.paragraph_format.left_indent = Cm(0.4)
    return paragraph


def _docx_table(doc, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            _set_docx_font(run, bold=True, size=10)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index + 1].cells[col_index]
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                _set_docx_font(run, size=9)
    doc.add_paragraph()


def _docx_code_block(doc, title: str, content: Any):
    _docx_paragraph(doc, title, bold=True)
    _docx_paragraph(doc, _pretty_json_text(content), monospace=True)


def _append_report_word(doc, report: RunOut) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("接口自动化测试报告")
    _set_docx_font(run, size=20, bold=True)

    _docx_paragraph(doc, f"报告时间：{_format_report_time(report.started_at)}    套件：{report.target_name or '-'}", bold=True)
    _docx_heading(doc, "报告摘要", level=2)
    _docx_table(doc, ["项目", "内容"], _report_summary_rows(report))

    _docx_heading(doc, "用例执行明细", level=2)
    for index, step in enumerate(report.steps, start=1):
        status = STATUS_LABELS.get(step.status, step.status)
        _docx_heading(doc, f"{index}. {step.case_name}（{status}）", level=3)
        _docx_paragraph(
            doc,
            f"方法：{step.method}    耗时：{_format_duration(step.duration_ms)}    "
            f"响应状态：{step.response_status if step.response_status is not None else '-'}",
        )
        _docx_paragraph(doc, f"请求 URL：{step.url}")

        if step.error_message:
            _docx_paragraph(doc, f"错误信息：{step.error_message}", bold=True)

        _docx_code_block(doc, "请求头", step.request_headers)
        _docx_code_block(doc, "请求体", step.request_body)
        _docx_code_block(doc, "响应头", step.response_headers)
        _docx_code_block(doc, "响应体", step.response_body)

        assertion_rows = _assertion_table_rows(step)
        if assertion_rows:
            _docx_paragraph(doc, "断言结果", bold=True)
            _docx_table(doc, ["类型", "说明", "期望", "实际", "结果"], assertion_rows)


def build_run_export_word(reports: List[RunOut]) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for index, report in enumerate(reports):
        if index > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _append_report_word(doc, report)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

