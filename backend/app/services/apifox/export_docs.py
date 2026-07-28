"""Apifox → 接口文档导出（HTML / Markdown / Word）。

复用 export_service 的范围过滤与 spec 解析，先归一化为「接口视图模型」，
再由三个渲染器分别产出对应格式；Word 用 python-docx。
"""

import io
from dataclasses import dataclass, field
from html import escape
from typing import Dict, List, Optional

from sqlalchemy.orm import Session

from app.repositories.apifox import schema_repo
from app.services.apifox.export_service import ExportOptions, _collect, _rows, _spec


@dataclass
class _Param:
    name: str
    location: str
    type: str
    required: bool
    desc: str


@dataclass
class _ApiView:
    method: str
    path: str
    name: str
    description: str
    params: List[_Param] = field(default_factory=list)
    body_type: str = "none"
    body_raw: str = ""
    response_schema: str = ""


_LOC_CN = {"query": "Query", "path": "Path", "header": "Header"}


def _params(spec: dict) -> List[_Param]:
    out: List[_Param] = []
    for loc, key in (("query", "query"), ("path", "path_params"), ("header", "headers")):
        for r in _rows(spec.get(key)):
            out.append(
                _Param(
                    name=str(r.get("key")).strip(),
                    location=_LOC_CN.get(loc, loc),
                    type=str(r.get("type") or "string"),
                    required=loc == "path" or bool(r.get("required")),
                    desc=str(r.get("desc") or ""),
                )
            )
    return out


def _views(db: Session, project_id: int, opts: ExportOptions):
    proj, endpoints, folders, _ = _collect(db, project_id, opts)
    groups: Dict[str, List[_ApiView]] = {}
    for ep in endpoints:
        spec = _spec(ep)
        body = spec.get("body") or {}
        resp_name = ""
        if ep.response_schema_id:
            sch = schema_repo.get_schema(db, ep.response_schema_id)
            resp_name = sch.name if sch else ""
        view = _ApiView(
            method=ep.method.upper(),
            path=ep.path or "/",
            name=ep.name or ep.path or "接口",
            description=ep.description or "",
            params=_params(spec),
            body_type=str(body.get("type") or "none"),
            body_raw=str(body.get("raw") or ""),
            response_schema=resp_name,
        )
        groups.setdefault(folders.get(ep.folder_id) or "未分组", []).append(view)
    return (proj.name if proj else "接口文档"), groups


# ---------- Markdown ----------
def build_markdown(db: Session, project_id: int, opts: ExportOptions) -> str:
    title, groups = _views(db, project_id, opts)
    lines: List[str] = [f"# {title}", ""]
    for folder, apis in groups.items():
        lines.append(f"## {folder}")
        lines.append("")
        for a in apis:
            lines.append(f"### {a.name}")
            lines.append("")
            lines.append(f"`{a.method}` `{a.path}`")
            lines.append("")
            if a.description:
                lines.append(a.description)
                lines.append("")
            if a.params:
                lines.append("| 参数 | 位置 | 类型 | 必填 | 说明 |")
                lines.append("| --- | --- | --- | --- | --- |")
                for p in a.params:
                    lines.append(f"| {p.name} | {p.location} | {p.type} | {'是' if p.required else '否'} | {p.desc} |")
                lines.append("")
            if a.body_type not in ("none", "") and a.body_raw:
                lines.append(f"请求体（{a.body_type}）：")
                lines.append("")
                lines.append("```")
                lines.append(a.body_raw)
                lines.append("```")
                lines.append("")
            if a.response_schema:
                lines.append(f"响应数据模型：`{a.response_schema}`")
                lines.append("")
    return "\n".join(lines)


# ---------- HTML ----------
_HTML_STYLE = (
    "body{font-family:-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;"
    "max-width:960px;margin:0 auto;padding:24px;color:#1f2329}"
    "h1{border-bottom:2px solid #e5e6eb;padding-bottom:8px}"
    "h3{margin-top:24px}"
    "table{border-collapse:collapse;width:100%;margin:8px 0}"
    "th,td{border:1px solid #e5e6eb;padding:6px 10px;text-align:left;font-size:14px}"
    "th{background:#f2f3f5}"
    "code{background:#f2f3f5;padding:2px 6px;border-radius:4px}"
    "pre{background:#f7f8fa;padding:12px;border-radius:6px;overflow:auto}"
    ".m{font-weight:700;color:#165dff}"
)


def build_html(db: Session, project_id: int, opts: ExportOptions) -> str:
    title, groups = _views(db, project_id, opts)
    out: List[str] = [
        "<!DOCTYPE html><html lang='zh'><head><meta charset='utf-8'>",
        f"<title>{escape(title)}</title><style>{_HTML_STYLE}</style></head><body>",
        f"<h1>{escape(title)}</h1>",
    ]
    for folder, apis in groups.items():
        out.append(f"<h2>{escape(folder)}</h2>")
        for a in apis:
            out.append(f"<h3>{escape(a.name)}</h3>")
            out.append(f"<p><span class='m'>{escape(a.method)}</span> <code>{escape(a.path)}</code></p>")
            if a.description:
                out.append(f"<p>{escape(a.description)}</p>")
            if a.params:
                out.append("<table><tr><th>参数</th><th>位置</th><th>类型</th><th>必填</th><th>说明</th></tr>")
                for p in a.params:
                    out.append(
                        f"<tr><td>{escape(p.name)}</td><td>{escape(p.location)}</td>"
                        f"<td>{escape(p.type)}</td><td>{'是' if p.required else '否'}</td>"
                        f"<td>{escape(p.desc)}</td></tr>"
                    )
                out.append("</table>")
            if a.body_type not in ("none", "") and a.body_raw:
                out.append(f"<p>请求体（{escape(a.body_type)}）：</p><pre>{escape(a.body_raw)}</pre>")
            if a.response_schema:
                out.append(f"<p>响应数据模型：<code>{escape(a.response_schema)}</code></p>")
    out.append("</body></html>")
    return "".join(out)


# ---------- Word (.docx) ----------
def build_docx(db: Session, project_id: int, opts: ExportOptions) -> bytes:
    from docx import Document  # 延迟导入，仅 Word 导出时加载

    title, groups = _views(db, project_id, opts)
    doc = Document()
    doc.add_heading(title, level=0)
    for folder, apis in groups.items():
        doc.add_heading(folder, level=1)
        for a in apis:
            doc.add_heading(a.name, level=2)
            doc.add_paragraph(f"{a.method} {a.path}")
            if a.description:
                doc.add_paragraph(a.description)
            if a.params:
                table = doc.add_table(rows=1, cols=5)
                table.style = "Light Grid Accent 1"
                for cell, text in zip(table.rows[0].cells, ("参数", "位置", "类型", "必填", "说明")):
                    cell.text = text
                for p in a.params:
                    cells = table.add_row().cells
                    cells[0].text = p.name
                    cells[1].text = p.location
                    cells[2].text = p.type
                    cells[3].text = "是" if p.required else "否"
                    cells[4].text = p.desc
            if a.body_type not in ("none", "") and a.body_raw:
                doc.add_paragraph(f"请求体（{a.body_type}）：")
                doc.add_paragraph(a.body_raw)
            if a.response_schema:
                doc.add_paragraph(f"响应数据模型：{a.response_schema}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- 出口 ----------
_DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def export_doc(db: Session, project_id: int, fmt: str, opts: Optional[ExportOptions] = None):
    """按格式返回 (内容, media_type, 文件名)；Word 返回 bytes，其余返回 str。"""
    opts = opts or ExportOptions()
    base = f"api-doc-project-{project_id}"
    if fmt == "markdown":
        return build_markdown(db, project_id, opts), "text/markdown; charset=utf-8", f"{base}.md"
    if fmt == "word":
        return build_docx(db, project_id, opts), _DOCX_MEDIA, f"{base}.docx"
    return build_html(db, project_id, opts), "text/html; charset=utf-8", f"{base}.html"
