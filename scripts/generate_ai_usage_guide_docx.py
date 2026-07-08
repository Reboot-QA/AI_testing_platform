#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate AI usage guidelines Word document for the AI testing platform."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
# Unicode escapes avoid Windows console encoding issues when saving.
OUTPUT = ROOT / "docs" / "AI\u4f7f\u7528\u89c4\u8303_V1.0.docx"


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=22, bold=True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=12, color=(100, 100, 100))
    p.paragraph_format.space_after = Pt(18)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 14, bold=True)
    return h


def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.line_spacing = 1.25


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.line_spacing = 1.25


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, cell_text in enumerate(row):
            run = table.rows[row_idx].cells[col_idx].paragraphs[0].add_run(cell_text)
            set_run_font(run, size=10)
    doc.add_paragraph()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    today = date.today().strftime("%Y年%m月%d日")
    add_title(doc, "AI 质量平台 · AI 使用规范")
    add_subtitle(doc, f"版本：V1.0    日期：{today}")

    add_para(
        doc,
        "本文档面向测试人员、项目管理员与系统管理员，说明 AI 质量平台中各类 AI 能力的用途、"
        "配置方式、使用流程及安全注意事项，帮助团队规范、高效地使用 AI 功能。",
    )

    # 1
    add_heading(doc, "1. 文档范围与适用对象", 1)
    add_bullets(
        doc,
        [
            "适用对象：测试工程师、测试负责人、部门管理员、平台管理员。",
            "适用场景：需求分析、用例设计、接口自动化、平台操作引导与演示。",
            "不适用：与本平台无关的通用问答、未授权的外部系统操作。",
        ],
    )

    # 2
    add_heading(doc, "2. AI 能力概览", 1)
    add_table(
        doc,
        ["功能模块", "入口路径", "主要能力", "是否调用大模型"],
        [
            ["AI 助手", "右下角悬浮按钮 / 右侧对话面板", "平台问答、浏览器自动演示", "Mock 或 LLM"],
            ["AI 智能生成", "用例管理 → AI 生成", "根据需求批量生成功能用例", "Mock 或 LLM"],
            ["需求文档解析", "需求文档", "上传 Word/PDF 后 AI 提取需求点", "Mock 或 LLM"],
            ["接口 AI 造数", "接口自动化 → 用例编辑", "按接口定义生成请求体/参数", "Mock 或 LLM"],
            ["助手自动化演示", "AI 助手欢迎区示例按钮", "一键演示项目管理等全流程", "预设脚本，无需 LLM"],
        ],
    )

    # 3
    add_heading(doc, "3. 大模型配置规范", 1)
    add_heading(doc, "3.1 配置入口与优先级", 2)
    add_numbered(
        doc,
        [
            "推荐在「系统管理 → 全局设置」中配置大模型 Provider（支持多模型、设默认、测试连接）。",
            "环境变量 backend/.env 中的 LLM_API_BASE、LLM_API_KEY、LLM_MODEL 为兜底配置。",
            "数据库中的全局设置优先级高于环境变量；生产环境应以全局设置为准。",
        ],
    )

    add_heading(doc, "3.2 Mock 模式与 LLM 模式", 2)
    add_bullets(
        doc,
        [
            "Mock 模式（默认）：不调用外部 API，使用本地模板生成内容，适合演示、联调、无密钥环境。",
            "LLM 模式：需配置有效的 API Base URL、API Key 与模型名称，并关闭 Mock 模式。",
            "未配置 API Key 或 Mock 模式开启时，所有 AI 生成类功能均走 Mock，界面会提示当前模式。",
            "常用免费模型示例：智谱 glm-4-flash、硅基流动 Qwen/Qwen2.5-7B-Instruct（以厂商文档为准）。",
        ],
    )

    add_heading(doc, "3.3 配置检查清单", 2)
    add_table(
        doc,
        ["检查项", "说明"],
        [
            ["Mock 开关", "演示环境可开启；正式生成用例前应关闭并验证连接"],
            ["API Key", "仅管理员配置；勿写入代码仓库、截图或聊天记录"],
            ["设为当前", "确保至少有一个 Provider 标记为「当前使用」"],
            ["测试连接", "保存前点击「测试连接」，确认模型可用"],
            ["批量参数", "AI 生成批次与并发由服务端配置（默认 batch=8, concurrency=4）"],
        ],
    )

    # 4
    add_heading(doc, "4. AI 助手使用规范", 1)
    add_heading(doc, "4.1 基本操作", 2)
    add_bullets(
        doc,
        [
            "点击右下角「AI 助手」悬浮按钮打开对话面板；可拖动悬浮按钮调整位置。",
            "输入自然语言问题，助手将流式回复操作步骤或平台说明。",
            "面板底部提示：支持浏览器自动操作；写入类操作需用户确认后执行。",
            "对话内容不持久化保存；退出登录后自动清空；重新登录后显示欢迎页。",
        ],
    )

    add_heading(doc, "4.2 一键演示（欢迎区示例）", 2)
    add_para(doc, "欢迎页提供四个全流程演示按钮，点击后自动在浏览器中执行，无需再次确认：")
    add_numbered(
        doc,
        [
            "帮我演示项目管理全流程 — 创建演示项目",
            "帮我演示需求管理全流程 — 选择项目并创建演示需求",
            "帮我演示用例管理全流程 — AI 生成用例并进入用例库",
            "帮我演示接口自动化管理全流程 — 环境、套件、演示用例与执行入口",
        ],
    )
    add_para(doc, "演示完成后约 1 秒对话区恢复欢迎状态，便于连续体验。")

    add_heading(doc, "4.3 自定义自动化指令", 2)
    add_bullets(
        doc,
        [
            "可使用「演示」「帮我创建」「自动导入 Swagger」等表述触发操作规划。",
            "助手返回「待执行操作」步骤列表后，请点击「确认执行」；可点击「取消」放弃。",
            "写入数据（创建项目、保存用例等）均须用户确认，只读导航类步骤按规划自动执行。",
            "Swagger 导入：文档地址以 /docs 结尾时，系统会自动转换为 /openapi.json。",
            "若某步失败，查看执行日志中的错误提示，确认当前页面与项目上下文是否正确。",
        ],
    )

    add_heading(doc, "4.4 提问建议", 2)
    add_table(
        doc,
        ["推荐做法", "不推荐做法"],
        [
            ["「如何在需求点中评审通过后再生成用例？」", "过于笼统的「怎么用 AI？」"],
            ["「帮我演示创建一个项目叫登录模块测试」", "要求助手操作与本平台无关的系统"],
            ["在当前业务页面提问（助手可感知页面路径）", "在自动化执行过程中频繁切换账号"],
        ],
    )

    # 5
    add_heading(doc, "5. AI 智能生成用例规范", 1)
    add_numbered(
        doc,
        [
            "在「需求点」中将待覆盖需求评审为「已评审」状态。",
            "进入「用例管理 → AI 智能生成」，选择目标项目与大模型 Provider。",
            "勾选已评审需求或填写需求描述，设置生成条数后点击「开始生成」。",
            "生成结果写入「用例库」，可在用例库中编辑、评审或导出 Excel。",
            "Mock 模式下生成内容为模板示例，正式项目验收前应切换 LLM 模式并人工复核。",
        ],
    )

    # 6
    add_heading(doc, "6. 需求与接口自动化中的 AI", 1)
    add_heading(doc, "6.1 需求文档 AI 提取", 2)
    add_bullets(
        doc,
        [
            "在「需求文档」上传 Word、PDF 等格式文件。",
            "使用 AI 提取需求点后，在「需求点」列表中维护、补充与评审。",
            "提取结果需人工核对，不可直接视为最终需求基线。",
        ],
    )

    add_heading(doc, "6.2 接口 AI 造数", 2)
    add_bullets(
        doc,
        [
            "在接口自动化用例编辑器中，根据接口定义一键生成请求体或参数。",
            "LLM 不可用时系统自动降级为 Mock 规则生成，并在结果中提示。",
            "涉及生产数据、隐私字段时，应脱敏后再使用 AI 造数。",
        ],
    )

    # 7
    add_heading(doc, "7. 数据安全与权限", 1)
    add_bullets(
        doc,
        [
            "部门数据隔离：同部门用户共享项目数据；管理员可查看全部部门。",
            "菜单与功能权限由「权限管理」分配，无权限用户无法访问对应 AI 入口。",
            "登录 Token 默认有效期 24 小时（ACCESS_TOKEN_EXPIRE_MINUTES=1440），过期需重新登录。",
            "API Key、数据库密码、JWT SECRET_KEY 不得提交至 Git 或对外传播。",
            "向大模型发送的需求/用例内容可能离开内网，敏感项目应使用私有化模型或 Mock 模式。",
            "退出登录时会清除 AI 助手对话及「AI 正在生成用例」等进行中的前端状态。",
        ],
    )

    # 8
    add_heading(doc, "8. 常见问题（FAQ）", 1)
    faq = [
        ("AI 生成一直是 Mock？", "检查全局设置中 Mock 模式是否关闭，且当前 Provider 已配置 Key 并通过测试。"),
        ("助手点击示例没反应？", "确认未处于「加载中」或「执行中」状态；刷新页面后重试。"),
        ("自动化找不到按钮？", "确保已选择正确项目；Swagger/套件操作建议使用助手预设演示或 invoke 类指令。"),
        ("生成用例质量不佳？", "补充清晰需求描述、减少单次条数、更换模型，并安排测试人员评审。"),
        ("演示后对话消失？", "设计如此：演示成功或退出登录后会重置为欢迎页，不保留历史记录。"),
    ]
    for q, a in faq:
        p = doc.add_paragraph()
        rq = p.add_run(f"Q：{q}\n")
        set_run_font(rq, bold=True)
        ra = p.add_run(f"A：{a}")
        set_run_font(ra)
        p.paragraph_format.space_after = Pt(8)

    # 9
    add_heading(doc, "9. 最佳实践", 1)
    add_numbered(
        doc,
        [
            "上线前在 Mock 模式完成流程验证，再切换 LLM 做正式生成。",
            "AI 输出必须经人工评审后再纳入测试基线或回归集。",
            "使用助手演示培训新人，使用自然语言问答解决日常操作疑问。",
            "定期轮换 API Key，管理员变更后立即在全局设置中更新。",
            "对接口自动化优先维护环境与套件结构，再使用 AI 造数提高效率。",
        ],
    )

    add_heading(doc, "10. 修订记录", 1)
    add_table(
        doc,
        ["版本", "日期", "说明", "作者"],
        [
            ["V1.0", today, "首版：整理平台 AI 能力、配置、助手规范与安全要求", "AI 质量平台团队"],
        ],
    )

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
