# -*- coding: utf-8 -*-
"""生成通用测试部门 AI Agent 规划 Word 文档"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            set_run_font(run, bold=True, size=10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=10)
    doc.add_paragraph()


def build_document(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("测试部门\nAI Agent 建设规划（通用版）")
    set_run_font(run, size=22, bold=True, color=RGBColor(0x1A, 0x56, 0xDB))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"版本：V1.0    日期：{date.today().strftime('%Y年%m月%d日')}")
    set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_page_break()

    add_heading(doc, "文档修订记录", 1)
    add_table(
        doc,
        ["版本", "日期", "修订内容", "作者"],
        [["V1.0", date.today().strftime("%Y-%m-%d"), "初稿：通用测试部门 AI Agent 建设规划", "测试部门"]],
    )

    add_heading(doc, "一、背景与建设动因", 1)
    add_para(
        doc,
        "软件交付节奏加快、系统复杂度上升，传统测试模式在需求理解、用例设计、回归执行、缺陷分析等环节"
        "高度依赖人工经验，存在效率瓶颈与质量波动。大语言模型（LLM）与 Agent 技术的成熟，"
        "为测试部门提供了从「AI 辅助工具」向「AI 自主协作」演进的可能。",
    )
    add_heading(doc, "1.1 当前典型痛点", 2)
    add_bullet(doc, "需求文档冗长、变更频繁，测试人员理解成本高、覆盖易遗漏")
    add_bullet(doc, "用例编写重复劳动多，评审周期长，难以快速响应迭代")
    add_bullet(doc, "接口/UI/性能/安全等多类型测试工具割裂，知识难以沉淀复用")
    add_bullet(doc, "回归范围判断依赖经验，全量回归成本高、选择性回归风险难控")
    add_bullet(doc, "缺陷根因分析、测试报告撰写占用大量非创造性时间")

    add_heading(doc, "1.2 建设目标", 2)
    add_bullet(doc, "建立测试领域 Multi-Agent 体系，覆盖「测什么 → 怎么测 → 测完怎么办」全链路")
    add_bullet(doc, "AI 生成内容必须经人工审批后方可入库/执行（Human-in-the-loop）")
    add_bullet(doc, "Agent 与现有测试管理、CI/CD、缺陷系统、接口/UI 工具集成，而非另起炉灶")
    add_bullet(doc, "可度量：效率提升、缺陷前置率、用例采纳率、自动化覆盖率等 KPI 可追踪")

    add_heading(doc, "1.3 建设原则", 2)
    add_table(
        doc,
        ["原则", "说明"],
        [
            ["人机协同", "Agent 负责提效与初稿，测试工程师负责决策与质量把关"],
            ["渐进落地", "从单点 Copilot 到单 Agent，再到多 Agent 编排，避免一步到位"],
            ["工具优先", "Agent 价值来自可调用的企业工具，而非 Prompt  alone"],
            ["安全合规", "数据不出域、权限最小化、操作可审计"],
            ["可替换模型", "不绑定单一 LLM 厂商，支持私有化与多模型路由"],
        ],
    )

    add_heading(doc, "二、AI Agent 与 Copilot 的区别", 1)
    add_table(
        doc,
        ["维度", "Copilot（辅助）", "Agent（自主协作）"],
        [
            ["交互方式", "用户逐步发起单次任务", "用户描述目标，Agent 自主分解多步执行"],
            ["上下文", "单次对话或表单输入", "项目、需求、历史用例、执行结果等持续记忆"],
            ["能力边界", "生成文本/代码建议", "调用工具：查库、写用例、跑脚本、提缺陷"],
            ["错误处理", "用户自行修正", "重试、降级、请求人工确认"],
            ["适用场景", "单条用例、单次数据构造", "端到端质量任务（如版本回归准备）"],
        ],
    )

    add_heading(doc, "三、测试 Agent 体系设计", 1)
    add_para(doc, "建议按测试生命周期划分 Agent 角色，由 Orchestrator（编排 Agent）统一调度。", bold=True)

    add_heading(doc, "3.1 Agent 角色定义", 2)
    add_table(
        doc,
        ["Agent", "核心职责", "典型输入", "典型输出", "关键工具"],
        [
            ["Requirement Agent", "需求解析、测试点识别、变更影响分析", "PRD、用户故事、变更 diff", "测试点清单、风险项", "文档解析、需求库、代码 diff"],
            ["Test Design Agent", "功能/场景/边界/异常用例设计", "测试点、业务规则", "用例草稿（步骤+预期）", "用例库、相似用例检索"],
            ["API Test Agent", "接口用例、Mock、断言、数据驱动", "Swagger/OpenAPI、抓包", "接口用例与测试数据", "Postman/平台 API、Mock 服务"],
            ["UI Test Agent", "页面流程与元素定位建议", "原型、URL、DOM 快照", "UI 脚本草稿", "Playwright/Selenium、录屏"],
            ["Data Agent", "测试数据构造、脱敏、边界数据集", "表结构、接口 Schema", "数据集、SQL/脚本", "数据库、Faker、脱敏规则"],
            ["Execution Agent", "执行调度、环境准备、结果采集", "用例集、版本号", "执行报告、失败明细", "CI、测试平台、容器"],
            ["Analysis Agent", "失败归因、缺陷草稿、回归建议", "日志、报告、变更范围", "根因摘要、缺陷描述", "日志平台、APM、Git"],
            ["Orchestrator Agent", "任务分解、Agent 路由、进度与审批", "自然语言任务描述", "任务计划与状态", "任务队列、审批流"],
        ],
    )

    add_heading(doc, "3.2 典型协作场景", 2)
    add_heading(doc, "场景 A：迭代测试准备", 3)
    add_para(doc, "输入：「为 v2.3 迭代准备功能测试」")
    add_para(doc, "流程：Requirement Agent 解析变更 → Test Design Agent 生成用例 → 人工评审 → 入库")
    add_heading(doc, "场景 B：接口回归", 3)
    add_para(doc, "输入：「对支付模块做接口冒烟」")
    add_para(doc, "流程：API Test Agent 拉取接口定义 → 生成用例与数据 → Execution Agent 执行 → Analysis Agent 出报告")
    add_heading(doc, "场景 C：线上问题复盘", 3)
    add_para(doc, "输入：「分析上周生产缺陷并补充回归用例」")
    add_para(doc, "流程：Analysis Agent 聚合缺陷 → Requirement Agent 映射测试点 → Test Design Agent 补用例")

    add_heading(doc, "四、总体架构", 1)
    add_heading(doc, "4.1 逻辑架构（五层）", 2)
    add_table(
        doc,
        ["层级", "组件", "说明"],
        [
            ["交互层", "Agent 工作台、IM 机器人、IDE 插件", "任务发起、进度展示、审批操作"],
            ["编排层", "Orchestrator、任务状态机、审批网关", "多 Agent 调度、并行与依赖管理"],
            ["Agent 运行时", "ReAct/Plan-Execute 循环、Memory、Prompt 管理", "单 Agent 推理与工具调用"],
            ["工具层", "Tools/MCP 适配器", "对接测试平台、Jira、Git、CI、数据库等"],
            ["数据层", "用例库、需求库、执行历史、向量索引、审计日志", "结构化数据 + RAG 知识库"],
        ],
    )

    add_heading(doc, "4.2 与现有系统集成", 2)
    add_bullet(doc, "测试管理：TestRail、禅道、自研平台 — 用例/需求 CRUD、评审状态")
    add_bullet(doc, "缺陷管理：Jira、Tapd — 创建缺陷、关联用例与版本")
    add_bullet(doc, "代码与 CI：GitLab/Jenkins — 变更范围、触发流水线、获取构建结果")
    add_bullet(doc, "接口/UI：Postman、Playwright、自研自动化平台 — 执行与报告回写")
    add_bullet(doc, "知识库：Confluence、Wiki — RAG 检索业务规则与历史案例")

    add_heading(doc, "五、技术选型建议", 1)
    add_table(
        doc,
        ["类别", "推荐方案", "备注"],
        [
            ["LLM 接入", "OpenAI 兼容 API / 私有化模型", "按任务路由：复杂推理用大模型，分类/摘要用小模型"],
            ["Agent 框架", "LangGraph / 自研轻量运行时", "需支持 streaming、cancel、工具注册"],
            ["工具协议", "Function Calling + MCP", "统一工具描述，便于扩展"],
            ["向量检索", "Milvus / pgvector / Elasticsearch", "项目内需求、用例、缺陷语义搜索"],
            ["任务队列", "Redis / RabbitMQ / Celery", "长任务异步、重试"],
            ["可观测", "步骤日志 + Token/耗时统计", "问题排查与成本管控"],
        ],
    )

    add_heading(doc, "六、分阶段实施路线图", 1)
    add_heading(doc, "阶段 0：准备期（2–4 周）", 2)
    add_bullet(doc, "梳理测试流程痛点，选定 2–3 个高 ROI 试点场景")
    add_bullet(doc, "确认 LLM 来源（公有云/私有化）、数据安全边界与审批流程")
    add_bullet(doc, "建立 Prompt 库、用例模板、质量评估标准（采纳率定义）")

    add_heading(doc, "阶段 1：Copilot 增强（4–6 周）", 2)
    add_bullet(doc, "需求摘要、用例生成、接口数据构造等单点 AI 能力上线")
    add_bullet(doc, "与测试管理平台集成，生成内容默认 draft 状态")
    add_bullet(doc, "验收：单场景效率提升 30%+，用例采纳率 ≥ 50%")

    add_heading(doc, "阶段 2：单 Agent 试点（6–8 周）", 2)
    add_bullet(doc, "落地 Test Design Agent 或 API Test Agent 之一，具备工具调用能力")
    add_bullet(doc, "Agent 工作台：任务创建、步骤流、人工审批节点")
    add_bullet(doc, "验收：端到端任务（如「从需求生成 20 条用例」）成功率 ≥ 80%")

    add_heading(doc, "阶段 3：Multi-Agent 编排（8–12 周）", 2)
    add_bullet(doc, "Orchestrator 上线，支持多 Agent 串联与并行")
    add_bullet(doc, "Execution Agent + Analysis Agent 接入 CI/报告")
    add_bullet(doc, "RAG 知识库：历史用例、缺陷、业务规则检索增强")

    add_heading(doc, "阶段 4：规模化与优化（持续）", 2)
    add_bullet(doc, "预置任务模板：版本回归、冒烟、安全扫描清单、发布检查")
    add_bullet(doc, "质量反馈闭环：采纳率、漏测率驱动 Prompt/策略迭代")
    add_bullet(doc, "成本优化：缓存、小模型路由、批处理")

    add_heading(doc, "七、组织与治理", 1)
    add_heading(doc, "7.1 角色分工", 2)
    add_table(
        doc,
        ["角色", "职责"],
        [
            ["测试负责人", "目标与 KPI、试点选型、审批策略"],
            ["测试工程师", "Agent 任务发起、产出评审、例外处理"],
            ["测试开发/平台", "工具适配、Agent 运行时、集成开发"],
            ["安全/合规", "数据分级、脱敏、审计要求"],
            ["AI 运营（可兼职）", "Prompt 维护、模型路由、效果评估"],
        ],
    )

    add_heading(doc, "7.2 治理机制", 2)
    add_bullet(doc, "所有 AI 生成用例/脚本入库前必须人工评审（可配置自动通过规则仅限低风险场景）")
    add_bullet(doc, "生产环境执行、批量删除等高风险操作需二次确认")
    add_bullet(doc, "每月 Agent 效果复盘：采纳率、失败任务、Token 成本、用户反馈")
    add_bullet(doc, "Prompt 与工具变更走版本管理，支持回滚")

    add_heading(doc, "八、安全与合规", 1)
    add_bullet(doc, "数据分级：需求/代码/生产数据禁止明文送入公有 LLM，需脱敏或私有化部署")
    add_bullet(doc, "权限：Agent 继承当前用户权限，禁止越权访问项目与系统")
    add_bullet(doc, "审计：记录任务 ID、调用的工具、输入输出摘要、操作人、时间")
    add_bullet(doc, "密钥：API Key 集中托管，不入 Agent 日志与 Prompt")

    add_heading(doc, "九、风险与应对", 1)
    add_table(
        doc,
        ["风险", "影响", "应对"],
        [
            ["模型幻觉导致错误用例", "漏测或误测", "结构化校验 + 人工评审 + 相似用例对照"],
            ["过度依赖 AI 削弱测试能力", "团队技能退化", "强调人机协同，AI 做初稿人做决策"],
            ["成本不可控", "预算超支", "Token 配额、任务级上限、模型分级路由"],
            ["工具调用失控", "误删数据、误执行", "最大步数、超时、高风险操作审批"],
            ["集成复杂度高", "项目延期", "阶段 1 只做 Copilot，工具集成从只读查询开始"],
        ],
    )

    add_heading(doc, "十、成功指标（KPI）", 1)
    add_table(
        doc,
        ["指标", "说明", "阶段 1 目标", "阶段 3 目标"],
        [
            ["用例设计效率", "同等质量下编写时间缩短比例", "≥ 30%", "≥ 50%"],
            ["AI 产出采纳率", "评审后入库占比", "≥ 50%", "≥ 70%"],
            ["需求测试点覆盖率", "AI 辅助后覆盖提升", "基线 +10%", "基线 +25%"],
            ["回归准备时间", "版本发布前测试准备工时", "缩短 20%", "缩短 40%"],
            ["Agent 任务成功率", "无需人工补救的比例", "—", "≥ 85%"],
            ["缺陷前置率", "测试阶段发现占比提升", "维持或提升", "提升 10%"],
        ],
    )

    add_heading(doc, "十一、预算与资源（参考）", 1)
    add_table(
        doc,
        ["项", "说明", "量级参考"],
        [
            ["LLM API 费用", "按 Token 与任务量", "试点期 5k–20k/月，视调用量"],
            ["向量库/算力", "私有化可选", "云托管或复用现有 infra"],
            ["人力", "平台开发 1–2 人 + 测试参与试点", "阶段 2 起投入"],
            ["培训", "Prompt 编写、Agent 使用规范", "1–2 次部门内训"],
        ],
    )

    add_heading(doc, "十二、总结与下一步", 1)
    add_para(
        doc,
        "测试部门 AI Agent 建设应遵循「人机协同、工具驱动、分阶段落地」路径："
        "先从 Copilot 单点提效验证价值，再引入单 Agent 工具调用，最后通过 Orchestrator 实现多 Agent 质量闭环。"
        "建议选定一个迭代周期内的「需求→用例」或「接口冒烟」作为首个试点，"
        "8 周内完成阶段 1–2 并产出可量化 KPI，再决定是否扩大编排与执行 Agent 投入。",
    )

    add_heading(doc, "附录：预置 Agent 任务模板示例", 1)
    add_table(
        doc,
        ["模板名称", "触发语示例", "涉及 Agent"],
        [
            ["迭代测试准备", "为 {版本} 准备功能测试用例", "Requirement + Test Design"],
            ["接口冒烟", "对 {模块} 执行接口冒烟测试", "API Test + Execution + Analysis"],
            ["需求变更影响", "分析 {需求ID} 变更影响的测试范围", "Requirement + Analysis"],
            ["缺陷补测", "根据 {缺陷单} 补充回归用例", "Analysis + Test Design"],
            ["发布检查", "生成 {版本} 发布前测试检查清单", "Orchestrator + 多 Agent"],
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    for out in (
        root / "QA_AI_Agent_Plan_Generic_V1.0.docx",
        root / "docs" / "QA_AI_Agent_Plan_Generic_V1.0.docx",
    ):
        build_document(out)
