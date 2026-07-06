# -*- coding: utf-8 -*-
"""生成 AI 测试平台 Agent 规划 Word 文档"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            set_run_font(run, bold=True, size=10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=10)
    doc.add_paragraph()


def build_document(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 测试平台\n智能 Agent 建设规划")
    set_run_font(run, size=22, bold=True, color=RGBColor(0x1A, 0x56, 0xDB))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"版本：V1.0    日期：{date.today().strftime('%Y年%m月%d日')}")
    set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_page_break()

    # 修订记录
    add_heading(doc, "文档修订记录", 1)
    add_table(
        doc,
        ["版本", "日期", "修订内容", "作者"],
        [
            ["V1.0", date.today().strftime("%Y-%m-%d"), "初稿：基于现有平台能力梳理 Agent 演进规划", "AI 测试平台团队"],
        ],
    )

    add_heading(doc, "一、背景与建设目标", 1)
    add_para(
        doc,
        "AI 测试平台（FastAPI + Vue 3）已具备需求解析、功能用例 AI 生成、接口调试与自动化执行等核心能力。"
        "当前 AI 能力以「单次 LLM 调用 + Prompt 模板」为主，属于 Copilot 辅助模式。"
        "下一阶段目标是演进为可感知上下文、可调用工具、可自主规划与执行的多 Agent 协作体系，"
        "实现从「人驱动 AI 生成」到「AI Agent 驱动质量闭环」的升级。",
    )
    add_heading(doc, "1.1 建设目标", 2)
    add_bullet(doc, "构建统一的 Agent 编排层，支撑需求→用例→执行→报告的全链路智能自动化")
    add_bullet(doc, "保留现有多 LLM 提供商配置能力，Agent 可按任务类型路由模型")
    add_bullet(doc, "Agent 行为可观测、可审计、可人工介入（Human-in-the-loop）")
    add_bullet(doc, "与现有项目管理、权限体系、接口自动化模块无缝集成")

    add_heading(doc, "1.2 不在本期范围", 2)
    add_bullet(doc, "替换现有手工测试执行 UI 与工作流")
    add_bullet(doc, "自研大模型训练与微调（继续使用 OpenAI 兼容 API）")
    add_bullet(doc, "跨租户 SaaS 多实例隔离（沿用现有单部署架构）")

    add_heading(doc, "二、现状分析", 1)
    add_heading(doc, "2.1 已有 AI 能力", 2)
    add_table(
        doc,
        ["模块", "实现位置", "能力说明", "模式"],
        [
            ["需求点提取", "ai_service.py / requirements 路由", "上传 Word/PDF 等文档，分块调用 LLM 提取结构化需求点", "单次/流式 LLM"],
            ["功能用例生成", "ai_service.py / testcases 路由", "按需求批量生成 P0-P3 用例，支持并发分批与 SSE 流式", "单次/流式 LLM"],
            ["接口测试数据生成", "api_data_ai_service.py", "根据接口定义生成 body/query/path 测试数据", "单次 LLM + Mock"],
            ["LLM 配置", "settings_service.py", "多提供商、默认模型、Mock 模式", "配置管理"],
        ],
    )

    add_heading(doc, "2.2 现有痛点", 2)
    add_bullet(doc, "各 AI 功能相互独立，无法串联成「需求评审通过后自动出用例、出接口脚本」的流水线")
    add_bullet(doc, "LLM 输出仅 JSON 解析，缺少工具调用、重试策略与中间状态持久化")
    add_bullet(doc, "接口自动化已有执行引擎，但未与 AI 联动生成断言、变量提取与套件编排")
    add_bullet(doc, "缺少 Agent 任务历史、Token 消耗与质量反馈闭环")

    add_heading(doc, "三、Agent 愿景与定位", 1)
    add_para(
        doc,
        "将平台 AI 能力从「Prompt 工程」升级为「Agent 工程」：每个 Agent 拥有明确角色、"
        "工具集（Tools）、记忆（Memory）与协作协议（Orchestration），在人工监督下完成测试质量任务。",
    )
    add_heading(doc, "3.1 核心 Agent 角色", 2)
    add_table(
        doc,
        ["Agent 名称", "职责", "主要工具", "产出物"],
        [
            ["Requirement Agent", "解析/补充/评审需求点", "文档解析、需求 CRUD、相似需求检索", "结构化需求清单"],
            ["TestCase Agent", "生成功能/接口/安全等用例", "用例库查询、去重、批量生成", "待评审用例"],
            ["API Agent", "接口用例与数据构造", "Swagger 导入、抓包解析、api_runner", "接口用例/套件/测试数据"],
            ["Execution Agent", "调度执行与结果分析", "手工执行、套件执行、报告生成", "执行记录与缺陷建议"],
            ["Orchestrator Agent", "任务分解与 Agent 调度", "任务队列、状态机、人工审批节点", "端到端质量任务"],
        ],
    )

    add_heading(doc, "3.2 演进路径对比", 2)
    add_table(
        doc,
        ["维度", "当前（Copilot）", "目标（Agent）"],
        [
            ["交互", "用户填表单触发单次生成", "自然语言任务 + 多步自主执行"],
            ["上下文", "单次 Prompt 内文本", "项目/需求/用例/执行历史 RAG"],
            ["错误处理", "失败即 HTTP 500", "重试、降级 Mock、人工确认"],
            ["协作", "无", "Orchestrator 编排多 Agent"],
        ],
    )

    add_heading(doc, "四、总体架构设计", 1)
    add_para(doc, "架构分为五层，自下而上：数据层 → 工具层 → Agent 运行时 → 编排层 → 交互层。", bold=True)

    add_heading(doc, "4.1 架构分层", 2)
    add_bullet(doc, "数据层：MySQL 现有模型 + Agent 任务表、步骤日志表、向量索引（可选 pgvector/Redis）")
    add_bullet(doc, "工具层：封装现有 Service 为 Agent Tools（requirements、testcases、api_automation、settings）")
    add_bullet(doc, "Agent 运行时：基于 ReAct / Function Calling 循环，支持 streaming 与 cancel")
    add_bullet(doc, "编排层：Orchestrator 负责任务 DAG、并行度、审批门禁（如用例入库前需评审）")
    add_bullet(doc, "交互层：新增「Agent 工作台」页面 + 保留现有 AIGenerate / 需求页入口")

    add_heading(doc, "4.2 关键数据流（示例：需求到用例闭环）", 2)
    add_para(doc, "1. 用户上传需求文档并发起 Agent 任务「从文档生成评审用例」")
    add_para(doc, "2. Orchestrator 分解为：Requirement Agent 提取 → TestCase Agent 生成 → 写入 draft 状态")
    add_para(doc, "3. 到达人工审批节点，测试负责人在 UI 批量评审")
    add_para(doc, "4. 审批通过后 Execution Agent 可选触发冒烟执行并汇总报告")

    add_heading(doc, "五、技术方案", 1)
    add_heading(doc, "5.1 后端新增模块（建议）", 2)
    add_table(
        doc,
        ["模块", "路径建议", "说明"],
        [
            ["agent_runtime", "backend/app/services/agent/", "LLM 工具调用循环、消息历史、超时控制"],
            ["agent_tools", "backend/app/services/agent/tools/", "Tool 定义与现有 Service 适配"],
            ["agent_orchestrator", "backend/app/services/agent/orchestrator.py", "任务状态机与多 Agent 路由"],
            ["agent_models", "backend/app/models/agent_task.py", "任务、步骤、工具调用日志"],
            ["agent_router", "backend/app/routers/agents.py", "REST + SSE 任务流 API"],
        ],
    )

    add_heading(doc, "5.2 LLM 集成策略", 2)
    add_bullet(doc, "继续复用 get_effective_llm_config()，Agent 任务可指定 provider_id")
    add_bullet(doc, "优先使用支持 Function Calling / Tools 的模型；不支持时降级为 JSON Schema 约束输出")
    add_bullet(doc, "长文档仍采用现有 split_document_chunks 分块 + Map-Reduce 汇总")
    add_bullet(doc, "Mock 模式用于 CI 与离线演示，Agent 工具调用走本地 Service 桩")

    add_heading(doc, "5.3 前端规划", 2)
    add_bullet(doc, "新增 views/AgentWorkbench.vue：任务创建、实时步骤流、工具调用详情、审批操作")
    add_bullet(doc, "现有 AIGenerate.vue 保留，逐步增加「升级为 Agent 任务」入口")
    add_bullet(doc, "系统设置扩展：Agent 默认模型、最大步数、并发任务数、审批策略")

    add_heading(doc, "5.4 工具（Tools）清单（一期）", 2)
    add_table(
        doc,
        ["Tool 名称", "对应现有能力", "Agent"],
        [
            ["extract_requirements", "stream_extract_requirements", "Requirement"],
            ["create_requirement", "requirements CRUD", "Requirement"],
            ["generate_testcases", "stream_generate_batches", "TestCase"],
            ["list_testcases", "testcases 查询", "TestCase"],
            ["generate_api_data", "generate_api_request_data", "API"],
            ["import_swagger", "api_swagger_service", "API"],
            ["run_api_suite", "api_runner_service", "Execution"],
            ["get_project_context", "projects + 统计", "Orchestrator"],
        ],
    )

    add_heading(doc, "六、实施路线图", 1)
    add_heading(doc, "阶段一：Agent 基础设施（4–6 周）", 2)
    add_bullet(doc, "Agent 任务/步骤数据模型与 API")
    add_bullet(doc, "Tool 适配层 + 单 Agent（TestCase Agent）PoC")
    add_bullet(doc, "SSE 任务流与前端工作台 MVP")
    add_bullet(doc, "验收：可通过自然语言触发「为某需求生成 N 条用例」并完成入库")

    add_heading(doc, "阶段二：多 Agent 编排（4–6 周）", 2)
    add_bullet(doc, "Orchestrator + Requirement Agent 串联")
    add_bullet(doc, "人工审批节点与任务暂停/恢复")
    add_bullet(doc, "任务历史、失败重试、Token/耗时统计")
    add_bullet(doc, "验收：文档上传 → 需求提取 → 用例生成 → 待评审 全自动")

    add_heading(doc, "阶段三：接口与执行 Agent（6–8 周）", 2)
    add_bullet(doc, "API Agent：Swagger/抓包 → 接口用例 + AI 测试数据")
    add_bullet(doc, "Execution Agent：套件执行 + 失败归因摘要")
    add_bullet(doc, "可选 RAG：项目内需求/用例语义检索增强上下文")
    add_bullet(doc, "验收：接口套件 AI 生成并可一键执行出报告")

    add_heading(doc, "阶段四：优化与规模化（持续）", 2)
    add_bullet(doc, "Agent 质量评估（用例采纳率、执行通过率）反馈 Prompt/策略")
    add_bullet(doc, "预置任务模板（回归、冒烟、安全扫描清单）")
    add_bullet(doc, "与 CI/CD、定时调度（现有 schedule_service）深度集成")

    add_heading(doc, "七、安全与合规", 1)
    add_bullet(doc, "Agent 仅调用当前用户有权限的项目与工具（复用 JWT + 项目 owner 校验）")
    add_bullet(doc, "API Key 不落库到 Agent 日志，工具调用参数脱敏（password/token 字段）")
    add_bullet(doc, "高风险操作（批量删除、执行生产环境套件）需显式用户确认")
    add_bullet(doc, "Agent 步骤全量审计日志，保留 90 天（可配置）")

    add_heading(doc, "八、风险与应对", 1)
    add_table(
        doc,
        ["风险", "影响", "应对措施"],
        [
            ["模型不稳定/幻觉", "用例质量下降", "结构化输出校验 + 人工评审门禁 + Mock 降级"],
            ["Token 成本过高", "运营成本", "分批并发控制、缓存相似需求结果、小模型路由"],
            ["工具调用死循环", "资源耗尽", "最大步数/超时/单任务 Token 上限"],
            ["与现有 UI 重复", "用户困惑", "渐进式入口，Copilot 与 Agent 并存"],
        ],
    )

    add_heading(doc, "九、成功指标（KPI）", 1)
    add_table(
        doc,
        ["指标", "基线（当前）", "阶段二目标", "阶段三目标"],
        [
            ["需求→用例人工步骤数", "5+ 步手动操作", "≤ 2 步（含审批）", "1 步发起 + 审批"],
            ["用例生成后采纳率", "待统计", "≥ 60%", "≥ 75%"],
            ["接口用例 AI 辅助覆盖率", "仅测试数据", "30% 用例 AI 草稿", "50%+"],
            ["Agent 任务成功率", "—", "≥ 85%", "≥ 92%"],
        ],
    )

    add_heading(doc, "十、总结", 1)
    add_para(
        doc,
        "本规划在充分复用 AI 测试平台现有 FastAPI 服务与 LLM 配置体系的前提下，"
        "以最小侵入方式引入 Agent 运行时与编排层，分四个阶段实现从 Copilot 到 Multi-Agent 质量闭环的演进。"
        "建议优先落地 TestCase Agent PoC 与 Agent 工作台，验证工具调用与 SSE 体验后，再扩展 Orchestrator 与接口/执行 Agent。",
    )

    add_heading(doc, "附录 A：与现有代码映射", 1)
    add_table(
        doc,
        ["现有文件", "Agent 演进关联"],
        [
            ["backend/app/services/ai_service.py", "TestCase / Requirement Agent 核心 LLM 逻辑"],
            ["backend/app/services/api_data_ai_service.py", "API Agent 测试数据工具"],
            ["backend/app/services/api_runner_service.py", "Execution Agent 执行工具"],
            ["backend/app/services/settings_service.py", "模型路由与 Mock 配置"],
            ["frontend/src/views/AIGenerate.vue", "逐步迁移为 Agent 任务入口"],
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    for out in (
        root / "AI_Agent_Plan_V1.0.docx",
        root / "docs" / "AI_Agent_Plan_V1.0.docx",
    ):
        build_document(out)
